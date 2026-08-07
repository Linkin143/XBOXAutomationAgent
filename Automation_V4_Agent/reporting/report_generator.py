"""Word document report generator — replaces Hcl.eDAT.WordReport.dll."""
from __future__ import annotations
import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..utils.logger import get_logger
from ..utils.helpers import timestamp_str

log = get_logger("report_generator")


@dataclass
class TestResultRow:
    game_name: str
    test_case: str
    status: str          # "PASS" | "FAIL"
    latency_ms: float = 0.0
    notes: str = ""
    timestamp: str = field(default_factory=lambda: datetime.now().strftime("%H:%M:%S"))


class ReportGenerator:
    """Accumulates test results and saves them to a .docx Word document."""

    def __init__(self, output_dir: str = "reports", template_path: Optional[str] = None):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_path = template_path
        self.results: list[TestResultRow] = []
        self._doc: Optional[Document] = None

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def add_result(self, game_name: str, test_case: str, status: str,
                   latency_ms: float = 0.0, notes: str = "") -> None:
        self.results.append(TestResultRow(game_name, test_case, status, latency_ms, notes))
        log.debug(f"Result added: {game_name}/{test_case} → {status}")

    def save(self, filename: Optional[str] = None) -> str:
        """Build and save the Word document. Returns the saved file path."""
        if filename:
            out_path = self.output_dir / filename
        else:
            out_path = self.output_dir / f"TestReport_{timestamp_str()}.docx"

        doc = Document(self.template_path) if self.template_path else Document()
        self._build_document(doc)
        doc.save(str(out_path))
        log.info(f"Report saved: {out_path}")
        return str(out_path)

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _build_document(self, doc: Document) -> None:
        # Title
        title = doc.add_heading("Xbox Automation V4 — Test Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Run meta
        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
                     f"Total: {len(self.results)}  |  "
                     f"PASS: {sum(1 for r in self.results if r.status == 'PASS')}  |  "
                     f"FAIL: {sum(1 for r in self.results if r.status == 'FAIL')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Results table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(["Game", "Test Case", "Status", "Latency (ms)", "Notes"]):
            hdr[i].text = col
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True

        for row in self.results:
            cells = table.add_row().cells
            cells[0].text = row.game_name
            cells[1].text = row.test_case
            cells[2].text = row.status
            cells[3].text = f"{row.latency_ms:.1f}"
            cells[4].text = row.notes
            # Colour status cell
            status_run = cells[2].paragraphs[0].runs[0]
            if row.status == "PASS":
                status_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            else:
                status_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Summary paragraph
        doc.add_paragraph()
        pass_rate = (sum(1 for r in self.results if r.status == "PASS") / len(self.results) * 100
                     if self.results else 0.0)
        doc.add_paragraph(f"Pass Rate: {pass_rate:.1f}%")
